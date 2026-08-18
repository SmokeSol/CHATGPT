#!/usr/bin/env python3
from __future__ import annotations
import argparse, hashlib, io, json, pathlib, zipfile

def sha256(path):
    h=hashlib.sha256()
    with open(path,'rb') as f:
        for b in iter(lambda:f.read(1024*1024),b''): h.update(b)
    return h.hexdigest()

def signature(path):
    with open(path,'rb') as f: return f.read(16).hex()

def inspect_stat(path, fmt):
    import pyreadstat
    kwargs={'metadataonly':True,'output_format':'dict'}
    if fmt=='dta': _, meta = pyreadstat.read_dta(path, **kwargs)
    elif fmt=='sav': _, meta = pyreadstat.read_sav(path, **kwargs)
    else: raise ValueError(fmt)
    names=list(meta.column_names); labels=list(meta.column_labels or [])
    value_labels=getattr(meta,'variable_value_labels',None) or {}
    cat={}
    for i,n in enumerate(names):
        cat[n]={'label':labels[i] if i<len(labels) else None,'value_labels':value_labels.get(n,{})}
    return {'columns':len(names),'rows':getattr(meta,'number_rows',None),'variables':cat}

def inspect_xlsx(path):
    from openpyxl import load_workbook
    raw=pathlib.Path(path).read_bytes()
    wb=load_workbook(io.BytesIO(raw),read_only=True,data_only=True)
    out={}
    for ws in wb.worksheets:
        rows=ws.iter_rows(values_only=True); first=[]
        for _ in range(8):
            try:first.append(next(rows))
            except StopIteration:break
        out[ws.title]={'max_row':ws.max_row,'max_column':ws.max_column,'first_rows':[[None if v is None else str(v) for v in r] for r in first]}
    return {'sheets':out}

def inspect_auto(path):
    raw=pathlib.Path(path).read_bytes()
    if raw[:2]==b'PK':
        with zipfile.ZipFile(io.BytesIO(raw)) as z:
            names=set(z.namelist())
            if 'word/document.xml' in names:
                from docx import Document
                doc=Document(io.BytesIO(raw))
                paras=[p.text.strip() for p in doc.paragraphs if p.text.strip()]
                tables=[]
                for t in doc.tables:
                    tables.append([[c.text.strip() for c in row.cells] for row in t.rows])
                return {'kind':'docx','paragraphs':paras,'tables':tables}
            if 'xl/workbook.xml' in names:
                x=inspect_xlsx(path); x['kind']='xlsx'; return x
            return {'kind':'zip','members':sorted(names)[:200]}
    return {'kind':'unknown','size_bytes':len(raw),'signature_hex':raw[:32].hex()}

def main():
    ap=argparse.ArgumentParser(); ap.add_argument('--manifest',required=True); ap.add_argument('--rawdir',required=True); ap.add_argument('--outdir',required=True); a=ap.parse_args()
    mani=json.load(open(a.manifest,encoding='utf-8')); raw=pathlib.Path(a.rawdir); out=pathlib.Path(a.outdir); out.mkdir(parents=True,exist_ok=True)
    audit=[]; catalog={}
    for s in mani['sources']:
        p=raw/(s['id']+'.bin'); rec={'id':s['id'],'url':s['url'],'role':s['role'],'exists':p.exists()}
        if not p.exists(): rec['status']='MISSING'; audit.append(rec); continue
        rec.update({'size_bytes':p.stat().st_size,'sha256':sha256(p),'signature_hex':signature(p)})
        try:
            if s['format'] in ('dta','sav'): ins=inspect_stat(str(p),s['format'])
            elif s['format']=='xlsx': ins=inspect_xlsx(str(p))
            else: ins=inspect_auto(str(p))
            catalog[s['id']]=ins; rec['status']='PARSED'
            if 'columns' in ins: rec['columns']=ins['columns']; rec['rows']=ins.get('rows')
            if 'sheets' in ins: rec['sheets']={k:{'max_row':v['max_row'],'max_column':v['max_column']} for k,v in ins['sheets'].items()}
            if 'kind' in ins: rec['kind']=ins['kind']
        except Exception as e:
            rec['status']='FETCHED_PARSE_FAILED'; rec['parse_error']=type(e).__name__+': '+str(e)
        audit.append(rec)
    stat_ids=['rgph2014_individual_dta','rgph2014_household_dta','encdm2014_household_sav','encdm2014_individual_sav','afrobarometer_morocco_r6','afrobarometer_morocco_r8']
    counts={k:(catalog.get(k,{}).get('columns',0) or 0) for k in stat_ids}
    parsed=[x['id'] for x in audit if x['status']=='PARSED']
    core_ok=all(counts[k]>0 for k in ['rgph2014_individual_dta','rgph2014_household_dta','afrobarometer_morocco_r6','afrobarometer_morocco_r8'])
    feasibility={'status':'SOURCE_CATALOG_READY' if core_ok else 'SOURCE_CATALOG_INCOMPLETE','raw_column_counts':counts,'parsed_source_count':len(parsed),'parsed_sources':parsed,'target_rich_dimensions_min':60,'important_note':'Raw column count is only a feasibility ceiling. Scientific feature selection must exclude leakage, redundancy, unsupported sensitive features and post-cutoff variables.'}
    json.dump({'manifest_id':mani['manifest_id'],'sources':audit},open(out/'source_fetch_audit.json','w'),ensure_ascii=False,indent=2)
    json.dump(catalog,open(out/'variable_catalog.json','w'),ensure_ascii=False,indent=2)
    json.dump(feasibility,open(out/'richness_feasibility.json','w'),ensure_ascii=False,indent=2)
    print(json.dumps(feasibility,indent=2))
if __name__=='__main__': main()
